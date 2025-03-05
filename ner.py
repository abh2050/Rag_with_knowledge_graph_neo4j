import docx
from typing import Tuple, List, Optional
from dataclasses import dataclass, field
import os
import re
import json
import argparse
from pathlib import Path
import logging
from tqdm import tqdm
import spacy
import signal
import sys  # Missing import needed for sys.exit()
from dataclasses import asdict  # Missing import needed for wandb config update
from PyPDF2 import PdfReader

from sklearn.feature_extraction.text import CountVectorizer
from bertopic import BERTopic
import torch
from transformers import AutoTokenizer, AutoModelForSequenceClassification

# Optional dependency
try:
    import wandb
    wandb_imported = True
except ImportError:
    wandb_imported = False

# Add this after the imports
import numpy as np

# Add these imports at the top of the file with other imports
from transformers import (
    AutoTokenizer,
    AutoModelForSequenceClassification,
    AutoModelForSeq2SeqLM,
    pipeline
)
from sentence_transformers import SentenceTransformer

class CustomJSONEncoder(json.JSONEncoder):
    """Custom JSON encoder to handle numpy types and other special objects."""
    def default(self, obj):
        if isinstance(obj, np.integer):
            return int(obj)
        if isinstance(obj, np.floating):
            return float(obj)
        if isinstance(obj, np.ndarray):
            return obj.tolist()
        if hasattr(obj, '__dict__'):
            return obj.__dict__
        return super().default(obj)


# --- Configuration ---
@dataclass
class ExtractorConfig:
    citation_linking: bool = True

@dataclass
class NLPConfig:
    ner_model: str = "en_core_web_sm"
    relation_extraction: bool = True
    relation_extraction_model: str = "distilbert-base-uncased"
    claim_detection: bool = True
    summarization: bool = True
    summarization_model: str = "facebook/bart-large-cnn"  # Changed from distilbert to BART
    embedding_model: str = "all-mpnet-base-v2"


@dataclass
class GraphConfig:
    uri: str
    user: str
    password: str

    def __post_init__(self):
        if not self.uri.startswith(('bolt://', 'neo4j://')):
            raise ValueError("Invalid Neo4j URI format")

@dataclass
class PipelineConfig:
    # Required parameters first
    graph: GraphConfig  # No default since it needs parameters
    
    # Optional parameters with defaults after
    output_dir: str = "./output"
    extractor: ExtractorConfig = field(default_factory=ExtractorConfig)
    nlp: NLPConfig = field(default_factory=NLPConfig)
    device: str = field(default_factory=lambda: "cuda" if torch.cuda.is_available() else "cpu")
    use_wandb: bool = False
    wandb_project: str = "scientific-document-pipeline"


# --- Data Classes ---
@dataclass
class Section:
    title: str
    text: str
    level: int
    position: int

@dataclass
class DocumentMetadata:
    id: str
    title: str
    authors: List[str]
    year: Optional[int] = None
    doi: Optional[str] = None
    publisher: Optional[str] = None

@dataclass
class Table:
    caption: str
    content: str  # Or a more structured representation
    position: int

@dataclass
class Figure:
    caption: str
    image_path: str  # Or a reference to the image data
    position: int

@dataclass
class Formula:
    text: str
    position: int

@dataclass
class Entity:
    text: str
    label: str
    start: int
    end: int
    confidence: float
    paragraph_id: Optional[str] = None
    section_id: Optional[str] = None
    normalized_form: Optional[str] = None
    linked_entity: Optional[str] = None  # e.g., a link to a database entry

@dataclass
class Relation:
    source: "Entity"  # Forward reference, needs quotes
    target: "Entity"
    relation_type: str
    confidence: float
    context: str

@dataclass
class Claim:
    text: str
    confidence: float
    is_hedged: bool
    supported_by: List[str]
    contradicted_by: List[str]

@dataclass
class Citation:
    id: str
    text: str
    reference_id: str
    context: str
    intent: str

@dataclass
class Topic:
    id: int
    keywords: List[str]
    weight: float

@dataclass
class ProcessedDocument:
    metadata: DocumentMetadata
    full_text: str
    sections: List[Section]
    tables: List[Table]
    figures: List[Figure]
    formulas: List[Formula]
    entities: List[Entity]
    relations: List[Relation]
    citations: List[Citation]
    topics: List[Topic]
    claims: List[Claim]
    summary: str
    embedding: List[float] #Store document embedding


# --- Custom Exceptions ---
class ExtractionError(Exception):
    pass

class GraphError(Exception):
    pass


# --- File Format Detection ---
def detect_file_format(file_path: str) -> str:
    """Detects the file format based on the file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return "pdf"
    elif ext == ".docx":
        return "docx"
    elif ext == ".xml":
        return "xml"
    elif ext == ".html":
        return "html"
    elif ext == ".txt":
        return "text"
    elif ext == ".json":
        return "json"
    else:
        return "unknown"


# --- Document Content Extraction ---
def extract_content_from_docx(file_path: str, config: ExtractorConfig) -> Tuple[str, List[Section], DocumentMetadata]:
    """Extract content from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])

        # Simple section extraction (improving upon this requires more complex heuristics)
        sections = []
        current_section_title = "Introduction"
        current_section_text = ""
        section_level = 1
        section_position = 0

        for para in doc.paragraphs:
          if para.style.name.startswith("Heading"):  # Heuristic: headings indicate new sections
              if current_section_text:
                  sections.append(Section(title=current_section_title, text=current_section_text, level=section_level, position=section_position))
                  section_position += 1
              current_section_title = para.text
              current_section_text = ""
              section_level = int(para.style.name[7:]) if para.style.name[7:].isdigit() else 1 # Extract heading level
          else:
              current_section_text += para.text + "\n"

        # Add the last section
        if current_section_text:
            sections.append(Section(title=current_section_title, text=current_section_text, level=section_level, position=section_position))

        # Add the full content as the first section
        sections.insert(0, Section(title="Full Content", text=full_text, level=1, position=0))

        # Attempt to extract metadata
        metadata = DocumentMetadata(
            id=os.path.basename(file_path),
            title=doc.core_properties.title or "Unknown Title",
            authors=[doc.core_properties.author] if doc.core_properties.author else []
            ,year=doc.core_properties.created.year if doc.core_properties.created else None
        )

        return full_text, sections, metadata
    except Exception as e:
        raise ExtractionError(f"Error extracting content from DOCX: {str(e)}")

def extract_content_from_pdf(file_path: str, config: ExtractorConfig) -> Tuple[str, List[Section], DocumentMetadata]:
    """Extract content from a PDF file."""
    try:
        # Initialize PDF reader
        reader = PdfReader(file_path)
        
        # Extract full text
        full_text = ""
        sections = []
        current_section = ""
        section_position = 0
        
        # Process each page
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            full_text += text + "\n"
            
            # Simple section detection based on common patterns
            potential_sections = re.split(r'\n(?=[0-9]+\.|[A-Z][a-z]+\s*\n)', text)
            
            for section_text in potential_sections:
                if len(section_text.strip()) > 0:
                    # Try to identify section title
                    lines = section_text.split('\n')
                    title = lines[0].strip()
                    content = '\n'.join(lines[1:]).strip()
                    
                    if len(content) > 0:  # Only add if there's content
                        sections.append(Section(
                            title=title,
                            text=content,
                            level=1,  # Default level
                            position=section_position
                        ))
                        section_position += 1
        
        # Add full content as first section
        sections.insert(0, Section(
            title="Full Content",
            text=full_text,
            level=0,
            position=0
        ))
        
        # Extract basic metadata
        metadata = DocumentMetadata(
            id=os.path.basename(file_path),
            title=os.path.splitext(os.path.basename(file_path))[0],
            authors=[],  # Would need more sophisticated extraction for authors
            year=None,  # Would need more sophisticated extraction for year
            doi=None,
            publisher=None
        )
        
        return full_text, sections, metadata
        
    except Exception as e:
        raise ExtractionError(f"Error extracting content from PDF: {str(e)}")

# Placeholder functions for other file types
def extract_content_from_xml(file_path: str, config: ExtractorConfig) -> Tuple[str, List[Section], DocumentMetadata]:
    raise NotImplementedError("XML extraction not yet implemented")

def extract_content_from_html(file_path: str, config: ExtractorConfig) -> Tuple[str, List[Section], DocumentMetadata]:
    raise NotImplementedError("HTML extraction not yet implemented")

def extract_content_from_text(file_path: str, config: ExtractorConfig) -> Tuple[str, List[Section], DocumentMetadata]:
    raise NotImplementedError("Text extraction not yet implemented")

def extract_content_from_json(file_path: str, config: ExtractorConfig) -> Tuple[str, List[Section], DocumentMetadata]:
    raise NotImplementedError("JSON extraction not yet implemented")


def extract_document_content(file_path: str, config: ExtractorConfig) -> Tuple[str, List[Section], List[Table], List[Figure], List[Formula], DocumentMetadata]:
    file_format = detect_file_format(file_path)
    empty_metadata = DocumentMetadata(id="", title="", authors=[])
    
    if file_format == "xml":
        text, sections, metadata = extract_content_from_xml(file_path, config)
        return text, sections, [], [], [], metadata or empty_metadata
    if file_format == "pdf":
        full_text, sections, metadata = extract_content_from_pdf(file_path, config)
        return full_text, sections, [], [], [], metadata
    elif file_format == "docx":
        full_text, sections, metadata = extract_content_from_docx(file_path, config)
        return full_text, sections, [], [], [], metadata
    elif file_format == "html":
         text, sections, metadata = extract_content_from_html(file_path, config), [], [], [], [], empty_metadata
         return text, sections, [], [], [], metadata
    elif file_format == "text":
        text, sections, metadata = extract_content_from_text(file_path, config), [], [], [], [], empty_metadata
        return text, sections, [], [], [], metadata
    elif file_format == "json":
         text, sections, metadata = extract_content_from_json(file_path, config), [], [], [], [], empty_metadata
         return text, sections, [], [], [], metadata
    else:
        raise ExtractionError(f"Unsupported file format: {file_format}")


# --- NLP Processing ---
class NLPProcessor:
    CACHE_DIR = "./models"  # Persistent cache directory

    def __init__(self, config: PipelineConfig):
        from transformers import AutoTokenizer, AutoModelForSequenceClassification, AutoModelForSeq2SeqLM, pipeline
        
        self.config = config
        self._nlp = None
        self._re_model = None
        self._re_tokenizer = None
        self._summarizer = None
        self._embedding_model = None
        self._topic_model = None  # Initialize topic model as None
        
        # Create cache directory if it doesn't exist
        os.makedirs(self.CACHE_DIR, exist_ok=True)
        
        try:
            logging.info("Loading NLP models...")
            
            # Check if models are already downloaded
            model_files = os.listdir(self.CACHE_DIR)
            
            # Relation extraction model
            if self.config.nlp.relation_extraction:
                model_name = self.config.nlp.relation_extraction_model.split('/')[-1]
                if not any(f.startswith(model_name) for f in model_files):
                    logging.info(f"Downloading relation extraction model {model_name}...")
                
                self._re_tokenizer = AutoTokenizer.from_pretrained(
                    self.config.nlp.relation_extraction_model,
                    cache_dir=self.CACHE_DIR,
                    use_fast=True,
                    local_files_only=True if model_files else False
                )
                self._re_model = AutoModelForSequenceClassification.from_pretrained(
                    self.config.nlp.relation_extraction_model,
                    num_labels=2,
                    cache_dir=self.CACHE_DIR,
                    local_files_only=True if model_files else False
                ).to(self.config.device)
                logging.info("Relation extraction model loaded")
            
            # Summarization model
            if self.config.nlp.summarization:
                # Skip loading heavy summarization model
                self._summarizer = None  # We'll use our custom summarization
                logging.info("Using lightweight summarization")
            
            # Embedding model
            self._embedding_model = SentenceTransformer(
                self.config.nlp.embedding_model,
                cache_folder=self.CACHE_DIR
            )
            logging.info("Embedding model loaded")
            
        except Exception as e:
            logging.error(f"Error initializing NLP models: {str(e)}")
            raise

    @property
    def nlp(self):
        """Lazy loading of spaCy model."""
        if self._nlp is None:
            try:
                self._nlp = spacy.load(self.config.nlp.ner_model)
            except OSError as e:
                logging.error(f"Failed to load spaCy model: {e}")
                logging.info("Try running: python -m spacy download en_core_web_sm")
                raise
        return self._nlp

    @nlp.setter
    def nlp(self, value):
        """Setter for nlp property."""
        self._nlp = value

    @property
    def re_tokenizer(self):
        """Lazy loading of relation extraction tokenizer."""
        if self._re_tokenizer is None:
            try:
                from transformers import AutoTokenizer
                self._re_tokenizer = AutoTokenizer.from_pretrained(
                    self.config.nlp.relation_extraction_model,
                    use_fast=True,
                    cache_dir="./models"  # Cache models locally
                )
            except Exception as e:
                logging.error(f"Failed to load relation extraction tokenizer: {e}")
                self._re_tokenizer = None
        return self._re_tokenizer

    @property
    def re_model(self):
        """Lazy loading of relation extraction model."""
        if self._re_model is None:
            try:
                from transformers import AutoModelForSequenceClassification
                self._re_model = AutoModelForSequenceClassification.from_pretrained(
                    self.config.nlp.relation_extraction_model,
                    num_labels=2,  # Binary classification: related or not
                    cache_dir="./models"  # Cache models locally
                ).to(self.config.device)
            except Exception as e:
                logging.error(f"Failed to load relation extraction model: {e}")
                self._re_model = None
        return self._re_model

    def perform_ner(self, text: str, chunk_size: int = 1000000) -> List[Entity]:
        """Process text in chunks to manage memory."""
        entities = []
        for i in range(0, len(text), chunk_size):
            chunk = text[i:i + chunk_size]
            doc = self.nlp(chunk)
            chunk_entities = [
                Entity(
                    text=ent.text,
                    label=ent.label_,
                    start=ent.start_char + i,  # Adjust for chunk offset
                    end=ent.end_char + i,
                    confidence=getattr(ent, 'confidence', 1.0),
                    paragraph_id=None,
                    section_id=None,
                    normalized_form=ent.lemma_,
                    linked_entity=None,
                )
                for ent in doc.ents
            ]
            entities.extend(chunk_entities)
        return entities

    def extract_relations_from_text(self, context: str, entity1: str, entity2: str) -> List[Relation]:
        """Extract relations between two entities from the given context."""
        try:
            if self.re_tokenizer is None or self.re_model is None:
                logging.warning(
                    f"Relation extraction model or tokenizer not available. "
                    f"Model: {self.config.nlp.relation_extraction_model}, "
                    f"Device: {self.config.device}"
                )
                return []

            inputs = self.re_tokenizer(
                f"{entity1} [SEP] {entity2} [SEP] {context}", 
                return_tensors="pt", 
                truncation=True, 
                padding=True
            ).to(self.config.device)
            
            with torch.no_grad():
                outputs = self.re_model(**inputs)
            logits = outputs.logits
            probs = torch.softmax(logits, dim=1)
            confidence = probs[0][1].item()  # Probability of being related
            
            if confidence > 0.5:
                return [Relation(
                    source=Entity(text=entity1, label="Unknown", start=0, end=0, confidence=1.0),
                    target=Entity(text=entity2, label="Unknown", start=0, end=0, confidence=1.0),
                    relation_type="related",
                    confidence=confidence,
                    context=context
                )]
            return []
        except Exception as e:
            logging.error(f"Error extracting relations: {str(e)}")
            return []

    def detect_claims(self, text: str) -> List[Claim]:  # Fixed syntax: List[Claim] instead of List<Claim
        """Detects claims in the given text."""
        hedging_words = ["suggest", "indicate", "may", "could", "possible", "likely"]
        sentences = re.split(r"(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s", text)  # More robust sentence splitting
        claims = []
        for sentence in sentences:
            if any(word in sentence.lower() for word in hedging_words):
                claims.append(Claim(
                    text=sentence,
                    confidence=0.5,
                    is_hedged=True,
                    supported_by=[],
                    contradicted_by=[]
                ))
        return claims

    def summarize_text(self, text: str) -> str:
        """Summarizes text using a hybrid approach."""
        try:
            # If text is short enough, return as is
            if len(text) < 1000:
                return text

            # Extractive summarization first
            sentences = text.split('.')
            important_sentences = []
            
            # Simple keyword-based importance scoring
            keywords = ['conclude', 'result', 'find', 'show', 'demonstrate', 'propose', 
                       'develop', 'analyze', 'study', 'observe', 'present', 'contribute',
                       'novel', 'new', 'important']
            
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                    
                # Score sentence based on keywords and length
                score = sum(1 for keyword in keywords if keyword in sentence.lower())
                if score > 0 and 10 < len(sentence.split()) < 50:
                    important_sentences.append(sentence)
            
            # Take top sentences (about 20% of original)
            num_sentences = max(3, min(10, len(sentences) // 5))
            summary = '. '.join(important_sentences[:num_sentences])
            
            if not summary:
                # Fallback to first few sentences if no important ones found
                summary = '. '.join(sentences[:5])
            
            return summary + '.'
            
        except Exception as e:
            logging.error(f"Error during summarization: {str(e)}")
            return text[:500] + "..."  # Fallback to simple truncation

    def generate_embeddings(self, text: str) -> List[float]:
        """Generates sentence embeddings for the given text."""
        try:
            if self._embedding_model is None:
                logging.warning("Embedding model not available")
                return [0.0] * 768  # Return zero vector as fallback
                
            embeddings = self._embedding_model.encode(text)
            return embeddings.tolist()
        except Exception as e:
            logging.error(f"Error generating embeddings: {e}")
            return [0.0] * 768  # Return zero vector as fallback

    def initialize_topic_model(self, all_texts: List[str]):
        """Initializes the BERTopic model with the given texts."""
        try:
            logging.info("Initializing BERTopic model...")
            vectorizer_model = CountVectorizer(stop_words="english")
            self._topic_model = BERTopic(
                vectorizer_model=vectorizer_model,
                language="english",
                calculate_probabilities=False,
                verbose=False,
                n_gram_range=(1, 2)
            )
            self._topic_model.fit(all_texts)  # Train on all texts
            logging.info("BERTopic model initialized.")
        except Exception as e:
            logging.error(f"Error initializing BERTopic model: {str(e)}")
            self._topic_model = None

    def extract_topics(self, text: str, num_topics: int = 5) -> List[Topic]:
        """Extracts topics from the given text using the pre-trained BERTopic model."""
        try:
            if self._topic_model is None:
                logging.warning("BERTopic model not initialized.")
                return []

            # Transform the document to get topic distribution
            topics, _ = self._topic_model.transform([text])
            topic_info = self._topic_model.get_topic_info()

            # Extract top N topics
            extracted_topics = []
            for topic_id in topics:
                if (topic_id == -1):  # Ignore outlier topics
                    continue
                keywords = self._topic_model.get_topic(topic_id)
                weight = topic_info[topic_info['Topic'] == topic_id]['Count'].values[0]
                extracted_topics.append(Topic(id=topic_id, keywords=[k[0] for k in keywords], weight=weight))

            return extracted_topics[:num_topics]  # Return top N topics

        except Exception as e:
            logging.error(f"Error extracting topics: {str(e)}")
            return []
# --- Citation Intent Classification ---
class CitationIntentClassifier:
    def __init__(self, device: str):
        self.device = device
        # More comprehensive keyword-based intent classification
        self.intent_keywords = {
            "background": ["background", "context", "previous", "related", "literature", "prior art", "historical"],
            "method": ["method", "approach", "technique", "procedure", "algorithm", "implementation", "simulation"],
            "result": ["result", "finding", "outcome", "conclusion", "observation", "performance", "evaluation"],
            "compare": ["compare", "contrast", "analyze", "evaluate", "benchmark", "assess"],
            "support": ["support", "validate", "confirm", "evidence", "verify", "corroborate"],
            "oppose": ["oppose", "contradict", "challenge", "disagree", "inconsistent", "limitation", "drawback"]
        }

    def classify_intent(self, context: str) -> Tuple[str, float]:
        """Classify the intent of a citation based on its context."""
        context_lower = context.lower()
        max_confidence = 0.0
        best_intent = "unknown"

        for intent, keywords in self.intent_keywords.items():
            keyword_count = sum(1 for keyword in keywords if keyword in context_lower)
            confidence = keyword_count / len(keywords)  # Normalize by keyword list length

            if confidence > max_confidence:
                max_confidence = confidence
                best_intent = intent

        return best_intent, max_confidence


# --- Knowledge Graph Interaction ---
from contextlib import contextmanager

class KnowledgeGraph:
    def __init__(self, config: GraphConfig):
        self.config = config
        self.driver = None

    @contextmanager
    def session_scope(self):
        """Context manager for Neo4j sessions."""
        if not self.driver:
            self._connect()
        session = self.driver.session()
        try:
            yield session
        finally:
            session.close()

    def _connect(self):
        from neo4j import GraphDatabase
        try:
            self.driver = GraphDatabase.driver(
                self.config.uri,  # Fix: use config.uri instead of self.uri
                auth=(self.config.user, self.config.password)
            )
            self.driver.verify_connectivity()
        except Exception as e:
            logging.error(f"Failed to connect to Neo4j: {e}")
            raise GraphError(f"Failed to connect to Neo4j: {e}")

    def add_document(self, document: "ProcessedDocument"):  # Forward reference
        """Adds a document and its metadata to the knowledge graph."""
        with self.session_scope() as session:
            session.run(
                """
                MERGE (d:Document {id: $doc_id})
                SET d.title = $title,
                    d.year = $year,
                    d.doi = $doi,
                    d.publisher = $publisher,
                    d.summary = $summary
                """,
                doc_id=document.metadata.id,
                title=document.metadata.title,
                year=document.metadata.year,
                doi=document.metadata.doi,
                publisher=document.metadata.publisher,
                summary=document.summary
            )
            # Add authors
            for author in document.metadata.authors:
                session.run(
                    """
                    MERGE (a:Author {name: $author_name})
                    MERGE (d:Document {id: $doc_id})-[:AUTHORED_BY]->(a)
                    """,
                    author_name=author,
                    doc_id=document.metadata.id
                )

    def add_entities(self, document_id: str, entities: List["Entity"]):
        """Adds entities to the knowledge graph and links them to the document."""
        with self.session_scope() as session:
            for entity in entities:
                session.run(
                    """
                    MERGE (e:Entity {text: $entity_text})
                    ON CREATE SET e.label = $entity_label, 
                                  e.confidence = $confidence
                    WITH e
                    MATCH (d:Document {id: $doc_id})
                    MERGE (d)-[:CONTAINS_ENTITY]->(e)
                    """,
                    entity_text=entity.text,
                    entity_label=entity.label,
                    confidence=entity.confidence,
                    doc_id=document_id
                )

    def add_relations(self, document_id: str, relations: List["Relation"]):  # Forward reference
        """Adds relations between entities to the knowledge graph."""
        with self.session_scope() as session:
            for relation in relations:
                session.run(
                    """
                    MATCH (s:Entity {name: $source_name})
                    MATCH (t:Entity {name: $target_name})
                    MERGE (s)-[r:RELATES_TO {type: $relation_type, confidence: $confidence}]->(t)
                    SET r.context = $context
                    """,
                    source_name=relation.source.text,
                    target_name=relation.target.text,
                    relation_type=relation.relation_type,
                    confidence=relation.confidence,
                    context=relation.context
                )

    def add_topics(self, document_id: str, topics: List["Topic"]):  # Forward reference
        """Adds topics to the knowledge graph and links them to the document."""
        with self.session_scope() as session:
            for topic in topics:
                session.run(
                    """
                    MERGE (t:Topic {id: $topic_id})
                    SET t.keywords = $keywords, t.weight = $weight
                    WITH t
                    MATCH (d:Document {id: $doc_id})
                    MERGE (d)-[:HAS_TOPIC]->(t)
                    """,
                    topic_id=topic.id,
                    keywords=topic.keywords,
                    weight=topic.weight,
                    doc_id=document_id
                )

    def add_citations(self, document_id: str, citations: List["Citation"]):  # Forward reference
        """Adds citations to the knowledge graph and links them to the document."""
        with self.session_scope() as session:
            for citation in citations:
                session.run(
                    """
                    MERGE (c:Citation {id: $citation_id})
                    SET c.text = $text, c.reference_id = $reference_id, c.intent = $intent
                    WITH c
                    MATCH (d:Document {id: $doc_id})
                    MERGE (d)-[:CITES]->(c)
                    """,
                    citation_id=citation.id,
                    text=citation.text,
                    reference_id=citation.reference_id,
                    intent=citation.intent,
                    doc_id=document_id
                )

    def add_claims(self, document_id: str, claims: List["Claim"]):  # Forward reference
        """Adds claims to the knowledge graph and links them to the document."""
        with self.session_scope() as session:
            for claim in claims:
                session.run(
                    """
                    MERGE (c:Claim {text: $text})
                    SET c.confidence = $confidence, c.is_hedged = $is_hedged
                    WITH c
                    MATCH (d:Document {id: $doc_id})
                    MERGE (d)-[:MAKES_CLAIM]->(c)
                    """,
                    text=claim.text,
                    confidence=claim.confidence,
                    is_hedged=claim.is_hedged,
                    doc_id=document_id
                )

    def close(self):
        """Closes the connection to the Neo4j database."""
        if self.driver:
            self.driver.close()
            logging.info("Disconnected from Neo4j.")


# --- Main Pipeline Class ---
class ScientificDocumentPipeline:
    """Orchestrates the entire scientific document analysis pipeline."""

    def __init__(self, config: PipelineConfig):
        self.config = config
        self.extractor_config = config.extractor
        self.nlp_config = config.nlp
        self.graph_config = config.graph

        self.nlp_processor = NLPProcessor(config)
        self.knowledge_graph = KnowledgeGraph(self.graph_config)
        self.citation_intent_classifier = CitationIntentClassifier(device=config.device)
        self.use_wandb = config.use_wandb

        # Remove explicit model loading since it's handled by the property
        if self.use_wandb and wandb_imported:
            wandb.init(project=config.wandb_project)
            wandb.config.update(asdict(config))

    def process_document(self, file_path: str) -> "ProcessedDocument":
        """Process a single document with optimized performance."""
        import time
        start_time = time.time()
        try:
            # Add file size check and logging
            file_size = os.path.getsize(file_path) / (1024 * 1024)  # MB
            logging.info(f"Processing file: {os.path.basename(file_path)} ({file_size:.1f}MB)")

            # 1. Extract Content
            logging.info("Extracting document content...")
            (
                full_text,
                sections,
                tables,
                figures,
                formulas,
                metadata,
            ) = extract_document_content(file_path, self.extractor_config)
            logging.info("Content extraction completed")

            # 2. NLP Processing with optimizations
            logging.info("Performing NER analysis...")
            entities = self.nlp_processor.perform_ner(full_text)
            logging.info(f"Found {len(entities)} entities")

            logging.info("Generating summary...")
            summary = self.nlp_processor.summarize_text(full_text)
            logging.info("Summary generated")

            logging.info("Generating embeddings...")
            embedding = self.nlp_processor.generate_embeddings(full_text)
            logging.info("Embeddings generated")

            # 3. Optimized Relation Extraction
            relations = []
            if self.nlp_config.relation_extraction and entities:
                logging.info("Extracting relations...")
                # Reduce the number of entity pairs to process
                max_entities = 400  # Limit the number of entities
                sampled_entities = entities[:max_entities]  # Take first N entities
                
                batch_size = 20  # Smaller batch size
                for i in range(0, len(sampled_entities), batch_size):
                    batch_entities = sampled_entities[i:i + batch_size]
                    for j, entity1 in enumerate(batch_entities):
                        # Only process nearby entities to reduce combinations
                        nearby_entities = batch_entities[j + 1:j + 6]  # Process only 5 nearby entities
                        for entity2 in nearby_entities:
                            context = full_text[max(0, entity1.start - 50):min(len(full_text), entity2.end + 50)]
                            extracted_relations = self.nlp_processor.extract_relations_from_text(
                                context, entity1.text, entity2.text
                            )
                            relations.extend(extracted_relations)
                    logging.info(f"Processed {min(i + batch_size, len(sampled_entities))}/{len(sampled_entities)} entities")
                logging.info(f"Found {len(relations)} relations")

            # 5. Topic Extraction
            logging.info("Extracting topics...")
            topics = self.nlp_processor.extract_topics(full_text)
            logging.info(f"Found {len(topics)} topics")

            # Create ProcessedDocument object
            processed_document = ProcessedDocument(
                metadata=metadata,
                full_text=full_text,
                sections=sections,
                tables=tables,
                figures=figures,
                formulas=formulas,
                entities=entities,
                relations=relations,
                citations=[],
                topics=topics,  # Add topics to the processed document
                claims=[],
                summary=summary,
                embedding=embedding,
            )

            processing_time = time.time() - start_time
            logging.info(f"Document processing completed in {processing_time:.2f} seconds")
            
            return processed_document

        except Exception as e:
            logging.error(f"Error processing document {file_path}: {str(e)}")
            raise ExtractionError(f"Failed to process document {file_path}: {str(e)}") from e

    def run(self, file_paths: List[str]):
        """Runs the pipeline with better progress tracking and skips processed files."""
        processed = 0
        failed = []
        all_texts = []

        total_files = len(file_paths)
        logging.info(f"Starting processing of {total_files} documents")

        # 1. Accumulate all texts, skipping already processed files
        files_to_process = []
        for idx, file_path in enumerate(file_paths, 1):
            output_path = Path(self.config.output_dir) / f"{os.path.basename(file_path)}.json"
            if output_path.exists():
                logging.info(f"Skipping {os.path.basename(file_path)} as it's already processed.")
                continue

            try:
                logging.info(f"Extracting text from document {idx}/{total_files}: {os.path.basename(file_path)}")
                full_text, _, _, _, _, _ = extract_document_content(file_path, self.extractor_config)
                all_texts.append(full_text)
                files_to_process.append(file_path)  # Keep track of files to process
            except Exception as e:
                logging.error(f"Failed to extract text from {file_path}: {str(e)}")
                failed.append((file_path, str(e)))
                continue

        # 2. Initialize topic model with all texts
        if all_texts:  # Only initialize if there are new texts
            logging.info("Initializing topic model with all documents...")
            self.nlp_processor.initialize_topic_model(all_texts)
            logging.info("Topic model initialized.")
        else:
            logging.info("No new documents to process, skipping topic model initialization.")

        # 3. Process each document
        for idx, file_path in enumerate(files_to_process, 1):
            try:
                logging.info(f"\nProcessing document {idx}/{len(files_to_process)}: {os.path.basename(file_path)}")
                processed_document = self.process_document(file_path)

                # Save to output
                output_path = Path(self.config.output_dir) / f"{os.path.basename(file_path)}.json"
                output_path.parent.mkdir(parents=True, exist_ok=True)

                with open(output_path, "w", encoding="utf-8") as f:
                    json.dump(asdict(processed_document), f, indent=4, ensure_ascii=False, cls=CustomJSONEncoder)

                logging.info(f"Saved processed document to {output_path}")
                
                # Load to Neo4j (add this part)
                neo4j_success = self.load_document_to_neo4j(processed_document)
                if neo4j_success:
                    logging.info(f"Document loaded to Neo4j: {os.path.basename(file_path)}")
                else:
                    logging.error(f"Failed to load document to Neo4j: {os.path.basename(file_path)}")
                
                processed += 1

            except Exception as e:
                failed.append((file_path, str(e)))
                logging.error(f"Failed to process {file_path}: {str(e)}")
                continue

        # Print summary
        logging.info(f"\nProcessing completed: {processed}/{total_files} documents successful")
        if failed:
            logging.error("\nFailed documents:")
            for path, error in failed:
                logging.error(f"- {path}: {error}")

        # Cleanup
        self.knowledge_graph.close()
        if self.use_wandb and wandb_imported:
            wandb.finish()

    def load_document_to_neo4j(self, processed_document: "ProcessedDocument"):
        """Load processed document into Neo4j knowledge graph."""
        try:
            logging.info("Loading document into Neo4j knowledge graph...")
            
            # Add document node
            self.knowledge_graph.add_document(processed_document)
            logging.info("Document node added")
            
            # Add entities and connect to document
            self.knowledge_graph.add_entities(processed_document.metadata.id, processed_document.entities)
            logging.info(f"Added {len(processed_document.entities)} entities")
            
            # Add relations
            self.knowledge_graph.add_relations(processed_document.metadata.id, processed_document.relations)
            logging.info(f"Added {len(processed_document.relations)} relations")
            
            # Add topics
            self.knowledge_graph.add_topics(processed_document.metadata.id, processed_document.topics)
            logging.info(f"Added {len(processed_document.topics)} topics")
            
            # Add claims
            self.knowledge_graph.add_claims(processed_document.metadata.id, processed_document.claims)
            logging.info(f"Added {len(processed_document.claims)} claims")
            
            # Add citations
            self.knowledge_graph.add_citations(processed_document.metadata.id, processed_document.citations)
            logging.info(f"Added {len(processed_document.citations)} citations")
            
            logging.info("Document successfully loaded into Neo4j knowledge graph")
            return True
            
        except Exception as e:
            logging.error(f"Error loading document into Neo4j: {str(e)}")
            return False


# --- Command Line Interface ---
def signal_handler(signum, frame):
    """Handle interrupt signals gracefully."""
    print("\nReceived interrupt signal. Cleaning up...")
    if 'pipeline' in globals():
        pipeline.knowledge_graph.close()
    if wandb_imported and wandb.run is not None:
        wandb.finish()
    sys.exit(0)

@contextmanager
def pipeline_scope(config: PipelineConfig):
    """Context manager for pipeline lifecycle."""
    pipeline = None
    try:
        pipeline = ScientificDocumentPipeline(config)
        yield pipeline
    finally:
        if pipeline:
            pipeline.knowledge_graph.close()
            if wandb_imported and wandb.run is not None:
                wandb.finish()

def main():
    # Load configuration
    with open('config.json', 'r') as f:
        config_data = json.load(f)
    
    # Define paths
    papers_dir = Path(config_data['papers_dir'])
    output_dir = Path(config_data['output_dir'])

    # Create pipeline configuration
    config = PipelineConfig(
        output_dir=str(output_dir),
        extractor=ExtractorConfig(),
        nlp=NLPConfig(),
        graph=GraphConfig(
            uri=config_data['neo4j']['uri'],
            user=config_data['neo4j']['user'],
            password=config_data['neo4j']['password']
        )
    )

    # Get all supported document files
    supported_extensions = [".pdf", ".docx", ".txt", ".xml", ".html", ".json"]
    file_paths = []
    for ext in supported_extensions:
        file_paths.extend(list(papers_dir.glob(f"*{ext}")))

    if not file_paths:
        print(f"No supported documents found in {papers_dir}")
        sys.exit(1)

    # Set up logging
    logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

    # Run the pipeline
    try:
        with pipeline_scope(config) as pipeline:
            pipeline.run([str(fp) for fp in file_paths])
    except Exception as e:
        logging.error(f"Pipeline failed: {str(e)}")
        sys.exit(1)

    print("Pipeline completed successfully.")

if __name__ == "__main__":
    main()